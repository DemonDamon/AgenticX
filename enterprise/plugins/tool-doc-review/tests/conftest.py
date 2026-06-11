#!/usr/bin/env python3
"""Pytest fixtures for tool-doc-review.

Author: Damon Li
"""

from __future__ import annotations

import sys
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

PLUGIN_DIR = Path(__file__).resolve().parent.parent
if str(PLUGIN_DIR) not in sys.path:
    sys.path.insert(0, str(PLUGIN_DIR))


@pytest.fixture
def clean_docx(tmp_path: Path) -> Path:
    path = tmp_path / "clean_sample.docx"
    document = Document()
    document.add_heading("Chapter One", level=1)
    body = document.add_paragraph("Body paragraph one.")
    body.runs[0].font.name = "Arial"
    body.runs[0].font.size = Pt(12)
    body2 = document.add_paragraph("Body paragraph two.")
    body2.runs[0].font.name = "Arial"
    body2.runs[0].font.size = Pt(12)
    document.add_paragraph("图1 Sample figure")
    document.add_paragraph("图2 Another figure")
    document.save(str(path))
    return path


@pytest.fixture
def flawed_docx(tmp_path: Path) -> Path:
    path = tmp_path / "flawed_sample.docx"
    document = Document()
    document.add_heading("Chapter One", level=1)
    document.add_heading("Skipped Level", level=3)
    normal_a = document.add_paragraph("Normal paragraph A")
    normal_a.style = document.styles["Normal"]
    normal_a.paragraph_format.space_before = Pt(0)
    normal_a.runs[0].font.name = "Arial"
    normal_a.runs[0].font.size = Pt(12)
    normal_b = document.add_paragraph("Normal paragraph B")
    normal_b.style = document.styles["Normal"]
    normal_b.paragraph_format.space_before = Pt(12)
    normal_b.runs[0].font.name = "Arial"
    normal_b.runs[0].font.size = Pt(14)
    document.add_paragraph("图1 First figure")
    document.add_paragraph("图3 Missing figure two")
    document.save(str(path))
    return path


@pytest.fixture
def flawed_labels(tmp_path: Path) -> Path:
    import json

    path = tmp_path / "flawed_labels.json"
    payload = {
        "expected": [
            {"category": "二类", "locator": "Heading level jump", "kind": "跳级"},
            {"category": "二类", "locator": "图2", "kind": "缺号"},
            {"category": "二类", "locator": "font-size-inconsistent", "kind": "字号"},
            {"category": "二类", "locator": "spacing-before-inconsistent", "kind": "段距"},
        ]
    }
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return path
